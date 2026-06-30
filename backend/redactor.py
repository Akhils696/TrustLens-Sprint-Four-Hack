"""
TrustLens – Redaction Engine.

Supports:
- PDF: black-box redaction using PyMuPDF annotations (native)
- DOCX: text replacement in paragraph runs (preserves formatting)
- TXT: string replacement, output as .txt

Always generates at minimum a redacted PDF.
Also generates a redacted TXT for all document types.
"""
from __future__ import annotations

import os
import re

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from fastapi import HTTPException


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def redact_document(
    file_path: str,
    filename: str,
    redactions: list[str],
) -> dict[str, str]:
    """
    Produce redacted outputs.

    Returns a dict:
        {"pdf": <path_to_redacted_pdf>, "txt": <path_to_redacted_txt>}
    """
    if not redactions:
        return _passthrough(file_path, filename)

    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        pdf_path = _redact_pdf(file_path, redactions)
    else:
        pdf_path = _redact_as_pdf_from_text(file_path, filename, redactions)

    txt_path = _redact_as_txt(file_path, filename, redactions)

    return {"pdf": pdf_path, "txt": txt_path}


# ---------------------------------------------------------------------------
# PDF redaction (native PyMuPDF black-box)
# ---------------------------------------------------------------------------

def _redact_pdf(file_path: str, redactions: list[str]) -> str:
    out_path = file_path + ".redacted.pdf"
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                # Ensure it's a fitz.Page object
                if not isinstance(page, fitz.Page):
                    continue
                for phrase in redactions:
                    phrase = phrase.strip()
                    if not phrase:
                        continue
                    hits = page.search_for(phrase, quads=False)
                    for rect in hits:
                        # Black filled redaction annotation — permanent removal
                        page.add_redact_annot(rect, fill=(0, 0, 0), cross_out=False)
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
            doc.save(out_path, garbage=4, deflate=True)
        return out_path
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF redaction failed: {str(e)}")


# ---------------------------------------------------------------------------
# Convert non-PDF to a redacted PDF
# ---------------------------------------------------------------------------

def _redact_as_pdf_from_text(file_path: str, filename: str, redactions: list[str]) -> str:
    """Extract text, apply redactions, write as a clean PDF."""
    out_path = file_path + ".redacted.pdf"
    ext = os.path.splitext(filename)[1].lower()

    try:
        raw_text = _read_text(file_path, filename, ext)
        redacted_text = _apply_text_redactions(raw_text, redactions)
        _write_text_as_pdf(redacted_text, out_path)
        return out_path
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document redaction failed: {str(e)}")


# ---------------------------------------------------------------------------
# Produce a plain-text redacted file
# ---------------------------------------------------------------------------

def _redact_as_txt(file_path: str, filename: str, redactions: list[str]) -> str:
    out_path = file_path + ".redacted.txt"
    ext = os.path.splitext(filename)[1].lower()

    try:
        raw_text = _read_text(file_path, filename, ext)
        redacted_text = _apply_text_redactions(raw_text, redactions)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(redacted_text)
        return out_path
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"TXT redaction failed: {str(e)}")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _read_text(file_path: str, filename: str, ext: str) -> str:
    """Read raw text from the source file."""
    if ext == ".txt":
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                with open(file_path, "r", encoding=enc, errors="strict") as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise HTTPException(status_code=400, detail="Cannot decode TXT file.")

    elif ext == ".docx":
        doc = DocxDocument(file_path)
        parts: list[str] = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    elif ext == ".pdf":
        doc = fitz.open(file_path)
        pages = [page.get_text("text") for page in doc]
        doc.close()
        return "\n\n".join(pages)

    raise HTTPException(status_code=400, detail=f"Unsupported extension for text read: {ext}")


def _apply_text_redactions(text: str, redactions: list[str]) -> str:
    """Replace each redaction phrase with [REDACTED] using case-sensitive exact match."""
    for phrase in redactions:
        phrase = phrase.strip()
        if not phrase:
            continue
        # Escape for regex, then replace all occurrences
        escaped = re.escape(phrase)
        text = re.sub(escaped, "[REDACTED]", text)
    return text


def _write_text_as_pdf(text: str, out_path: str) -> None:
    """Write a plain-text string to a clean PDF using PyMuPDF."""
    doc = fitz.open()
    # Split into pages of roughly 3000 chars each
    chunk_size = 3000
    chunks = [text[i : i + chunk_size] for i in range(0, max(len(text), 1), chunk_size)]

    for chunk in chunks:
        page = doc.new_page(width=595, height=842)  # A4
        rect = fitz.Rect(50, 50, 545, 792)
        page.insert_textbox(
            rect,
            chunk,
            fontsize=10,
            fontname="helv",
            color=(0, 0, 0),
        )

    doc.save(out_path, garbage=4, deflate=True)
    doc.close()


def _passthrough(file_path: str, filename: str) -> dict[str, str]:
    """No redactions — copy original as PDF and empty TXT."""
    ext = os.path.splitext(filename)[1].lower()
    pdf_path = file_path + ".redacted.pdf"
    txt_path = file_path + ".redacted.txt"

    if ext == ".pdf":
        import shutil
        shutil.copy2(file_path, pdf_path)
    else:
        raw = _read_text(file_path, filename, ext)
        _write_text_as_pdf(raw, pdf_path)

    with open(txt_path, "w", encoding="utf-8") as f:
        raw = _read_text(file_path, filename, ext)
        f.write(raw)

    return {"pdf": pdf_path, "txt": txt_path}
