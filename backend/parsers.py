"""
TrustLens – Document text extraction pipeline.

Supports: PDF (PyMuPDF), DOCX (python-docx), TXT (UTF-8/latin-1)
Returns: ExtractResult with text, page_count, word_count, char_count
Validates: file size (20 MB), encryption, corruption
"""
from __future__ import annotations

import os

import fitz  # PyMuPDF
from docx import Document
from fastapi import HTTPException

from models import ExtractResult

MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20 MB
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def validate_file(file_path: str, filename: str) -> None:
    """Validate file size and extension before extraction."""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: PDF, DOCX, TXT.",
        )

    size = os.path.getsize(file_path)
    if size > MAX_FILE_SIZE_BYTES:
        mb = size / (1024 * 1024)
        raise HTTPException(
            status_code=413,
            detail=f"File too large ({mb:.1f} MB). Maximum allowed size is 20 MB.",
        )

    if size == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")


def extract_text_from_file(file_path: str, filename: str) -> ExtractResult:
    """
    Extract text from PDF, DOCX, or TXT.
    Returns an ExtractResult with text content and metadata.
    """
    validate_file(file_path, filename)

    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        return _extract_txt(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".pdf":
        return _extract_pdf(file_path)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format.")


def _extract_txt(file_path: str) -> ExtractResult:
    """Extract text from a plain-text file, trying UTF-8 then latin-1."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding, errors="strict") as f:
                text = f.read()
            return _build_result(text, page_count=1)
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="Could not decode TXT file. Unsupported encoding.")


def _extract_docx(file_path: str) -> ExtractResult:
    """Extract text from a DOCX file preserving paragraph order."""
    try:
        doc = Document(file_path)
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs)
        return _build_result(text, page_count=1)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {str(e)}")


def _extract_pdf(file_path: str) -> ExtractResult:
    """Extract text from a PDF using PyMuPDF, with encryption and corruption checks."""
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to open PDF — file may be corrupted: {str(e)}")

    if doc.is_encrypted:
        doc.close()
        raise HTTPException(
            status_code=400,
            detail="PDF is password-protected. Please provide an unencrypted document.",
        )

    try:
        page_texts: list[str] = []
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")  # plain text mode
            if text.strip():
                # Prefix each page for context
                page_texts.append(f"[Page {page_num}]\n{text}")
        page_count = len(doc)
        doc.close()

        text = "\n\n".join(page_texts)
        return _build_result(text, page_count=page_count)
    except Exception as e:
        doc.close()
        raise HTTPException(status_code=500, detail=f"Failed to extract text from PDF: {str(e)}")


def _build_result(text: str, page_count: int) -> ExtractResult:
    words = text.split()
    return ExtractResult(
        text=text,
        page_count=page_count,
        word_count=len(words),
        char_count=len(text),
    )
